# backend/utils/docx_parser.py
from docx import Document
import json

def parse_docx(file_path: str) -> dict:
    """Extrae texto de DOCX y devuelve estructura básica"""
    doc = Document(file_path)
    
    # Estrategia de extracción
    text_blocks = []
    current_section = ""
    
    for para in doc.paragraphs:
        # Detectar secciones por estilo o palabras clave
        if para.style.name.startswith('Heading'):
            current_section = para.text
        else:
            text_blocks.append({
                "section": current_section,
                "content": para.text
            })
    
    # Extraer tablas si existen
    tables_data = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            table_data.append([cell.text for cell in row.cells])
        tables_data.append(table_data)
    
    return {
        "raw_text": "\n".join([p.text for p in doc.paragraphs]),
        "structured_blocks": text_blocks,
        "tables": tables_data
    }
